#!/usr/bin/env python3
"""
Converts the markdown IT Due Diligence document to DOCX format.
Preserves formatting (headers, tables, lists, bold/italic).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def parse_markdown_to_docx(md_path, docx_path):
    """Parse markdown and create DOCX with proper formatting."""

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    current_table = None
    table_headers = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines (but preserve them for spacing in some cases)
        if not line.strip():
            i += 1
            continue

        # Headers (# style)
        if line.startswith('# **'):
            # Main section header
            match = re.match(r'# \*\*(.+?)\*\*', line)
            if match:
                heading = match.group(1)
                p = doc.add_paragraph()
                run = p.add_run(heading)
                run.bold = True
                run.font.size = Pt(14)
                p.style = 'Heading 1'

        elif line.startswith('## **'):
            # Subsection header
            match = re.match(r'## \*\*(.+?)\*\*', line)
            if match:
                heading = match.group(1)
                p = doc.add_paragraph(heading, style='Heading 2')
                p.runs[0].bold = True

        # Markdown emphasis text (descriptive sections)
        elif line.startswith('*') and line.endswith('*') and not line.startswith('| '):
            # Italic descriptive text
            text = line.strip('*').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)

        # Tables
        elif line.startswith('| '):
            # Start or continue table
            if current_table is None:
                # Parse header row
                headers = [cell.strip() for cell in line.split('|')[1:-1]]
                current_table = doc.add_table(rows=1, cols=len(headers))
                current_table.style = 'Light Grid Accent 1'

                # Set header row
                header_cells = current_table.rows[0].cells
                for j, header in enumerate(headers):
                    header_cells[j].text = header
                    # Bold header
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                table_headers = headers
                i += 1

                # Skip separator line (|---|---|...)
                if i < len(lines) and re.match(r'^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|$', lines[i]):
                    i += 1
                continue

            elif current_table is not None:
                # Add data row
                cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells_data) == len(table_headers):
                    row_cells = current_table.add_row().cells
                    for j, cell_data in enumerate(cells_data):
                        # Parse cell content for bold/italic
                        cell_para = row_cells[j].paragraphs[0]
                        parse_inline_formatting(cell_para, cell_data)

        else:
            # Regular paragraph
            if current_table is not None:
                current_table = None

            if line.strip():
                # Parse inline formatting (bold, italic)
                p = doc.add_paragraph()
                parse_inline_formatting(p, line)

        i += 1

    doc.save(docx_path)
    print(f"✓ DOCX created: {docx_path}")


def parse_inline_formatting(paragraph, text):
    """Parse inline formatting: **bold**, *italic*, ***bold-italic***."""

    # Remove leading/trailing whitespace
    text = text.strip()

    # Pattern for ***text*** (bold-italic), **text** (bold), *text* (italic)
    pattern = r'(\*{1,3})(.+?)\1'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add normal text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        # Add formatted text
        formatted_text = match.group(2)
        markers = match.group(1)

        run = paragraph.add_run(formatted_text)
        if len(markers) == 3:
            run.bold = True
            run.italic = True
        elif len(markers) == 2:
            run.bold = True
        elif len(markers) == 1:
            run.italic = True

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


if __name__ == '__main__':
    md_file = '/Users/Imre/PythonProjects/python-for-ai/hikari-slides/doc/IT_Due_Diligence_LegalTech_Mikrovallalat.docx.md'
    docx_file = '/Users/Imre/PythonProjects/python-for-ai/hikari-slides/doc/IT_Due_Diligence_LegalTech_Mikrovallalat.docx'

    parse_markdown_to_docx(md_file, docx_file)
