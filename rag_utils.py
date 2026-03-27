"""
rag_utils.py — Retrieval-Augmented Generation helpers.

Provides:
  - get_collection()       : opens / creates the persistent ChromaDB collection
  - get_stats_collection() : opens / creates the index-stats ChromaDB collection
  - index_documents(dir)   : incrementally indexes files from the uploads folder
  - get_index_stats()      : returns per-file indexing statistics from ChromaDB
  - search_documents(query): semantic search; returns relevant chunks or None
  - search_web(query)      : DuckDuckGo fallback; returns formatted snippets or None
"""

import os
from datetime import datetime, timezone
import chromadb

# ── Paths ──────────────────────────────────────────────────────────────────────
# Both paths live at the project root (same directory as this file).
_ROOT       = os.path.dirname(os.path.abspath(__file__))
CHROMA_DIR  = os.path.join(_ROOT, "chroma_db")

COLLECTION_NAME       = "dropbox_docs"
STATS_COLLECTION_NAME = "index_stats"

# Chunking parameters: 500-character chunks with 50-character overlap so
# context is not lost at chunk boundaries.
CHUNK_SIZE    = 500
CHUNK_OVERLAP = 50

# Supported file extensions for text extraction.
SUPPORTED_EXTS = {".pdf", ".docx", ".xlsx", ".txt"}


# ── Collection accessors ───────────────────────────────────────────────────────
def get_collection() -> chromadb.Collection:
    """Return the persistent ChromaDB collection, creating it if necessary."""
    client = chromadb.PersistentClient(path=CHROMA_DIR)
    collection = client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
    )
    return collection


def get_stats_collection() -> chromadb.Collection:
    """Return the persistent index-stats ChromaDB collection (no embeddings needed)."""
    client = chromadb.PersistentClient(path=CHROMA_DIR)
    return client.get_or_create_collection(name=STATS_COLLECTION_NAME)


# ── Page counter ───────────────────────────────────────────────────────────────
def _count_pages(path: str, ext: str) -> int:
    """Estimate the number of pages / logical units in a document."""
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            return len(PdfReader(path).pages)
        if ext == ".docx":
            from docx import Document
            paragraphs = len(Document(path).paragraphs)
            return max(paragraphs // 10, 1)
        if ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(path, data_only=True, read_only=True)
            count = len(wb.worksheets)
            wb.close()
            return max(count, 1)
        if ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as fh:
                lines = sum(1 for _ in fh)
            return max(lines // 50, 1)
    except Exception:
        pass
    return 1


# ── Text extraction ────────────────────────────────────────────────────────────
def _extract_text(path: str) -> str:
    """Extract plain text from a PDF, DOCX, XLSX, or TXT file."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    if ext == ".xlsx":
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        lines = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                line = " ".join(str(c) for c in row if c is not None)
                if line.strip():
                    lines.append(line)
        return "\n".join(lines)

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as fh:
            return fh.read()

    return ""


# ── Text chunking ──────────────────────────────────────────────────────────────
def _chunk_text(text: str) -> list[str]:
    """Split text into overlapping fixed-size chunks."""
    chunks = []
    start = 0
    while start < len(text):
        end   = min(start + CHUNK_SIZE, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


# ── Document indexer ───────────────────────────────────────────────────────────
def index_documents(uploads_dir: str) -> None:
    """
    Incrementally index all supported files in uploads_dir into ChromaDB.

    - Skips files whose mtime has not changed since last index.
    - Re-indexes files that were modified.
    - Removes chunks for files that have been deleted.
    """
    if not os.path.isdir(uploads_dir):
        return

    collection       = get_collection()
    stats_collection = get_stats_collection()

    # Build a map of filename → mtime from what is already in ChromaDB.
    all_meta   = collection.get(include=["metadatas"])["metadatas"]
    indexed: dict[str, str] = {}   # filename -> stored mtime string
    for meta in all_meta:
        if meta and "filename" in meta and "mtime" in meta:
            indexed[meta["filename"]] = meta["mtime"]

    # Scan the uploads directory.
    disk_files: set[str] = set()
    for fname in os.listdir(uploads_dir):
        fpath = os.path.join(uploads_dir, fname)
        if not os.path.isfile(fpath):
            continue
        ext = os.path.splitext(fname)[1].lower()
        if ext not in SUPPORTED_EXTS:
            continue

        disk_files.add(fname)
        mtime = str(os.stat(fpath).st_mtime)

        # Skip unchanged files.
        if indexed.get(fname) == mtime:
            continue

        # Remove stale chunks before re-indexing.
        if fname in indexed:
            collection.delete(where={"filename": fname})

        # Extract, chunk, and upsert.
        text   = _extract_text(fpath)
        if not text.strip():
            continue
        chunks = _chunk_text(text)
        if not chunks:
            continue

        ids       = [f"{fname}_chunk_{i}" for i in range(len(chunks))]
        metadatas = [
            {"filename": fname, "mtime": mtime, "chunk_index": i}
            for i in range(len(chunks))
        ]
        collection.upsert(documents=chunks, ids=ids, metadatas=metadatas)

        # ── Persist indexing statistics ────────────────────────────────────────
        chars         = len(text)
        tokens_approx = chars // 4
        pages         = _count_pages(fpath, ext)
        indexed_at    = datetime.now(timezone.utc).isoformat()

        stats_collection.upsert(
            ids=[fname],
            documents=[fname],
            metadatas=[{
                "filename":     fname,
                "pages":        pages,
                "chunks":       len(chunks),
                "chars":        chars,
                "tokens_approx": tokens_approx,
                "indexed_at":   indexed_at,
            }],
        )

    # Remove chunks (and stats) for files that were deleted from the uploads folder.
    for fname in set(indexed.keys()) - disk_files:
        collection.delete(where={"filename": fname})
        try:
            stats_collection.delete(ids=[fname])
        except Exception:
            pass


# ── Index statistics ───────────────────────────────────────────────────────────
def get_index_stats() -> list[dict]:
    """
    Return per-file indexing statistics stored in the index_stats collection.

    Each dict contains: filename, pages, chunks, chars, tokens_approx, indexed_at.
    Returns an empty list if no files have been indexed yet.
    """
    try:
        stats = get_stats_collection().get(include=["metadatas"])["metadatas"]
        return [m for m in stats if m]
    except Exception:
        return []


# ── Semantic search ────────────────────────────────────────────────────────────
def search_documents(
    query: str,
    n_results: int = 4,
    threshold: float = 0.7,
) -> list[str] | None:
    """
    Query ChromaDB for chunks relevant to query.

    Returns a list of text chunks whose cosine distance is ≤ threshold,
    or None if the collection is empty or no chunk meets the threshold.

    Cosine distance in ChromaDB = 1 - cosine_similarity.
    Distance 0 → identical; distance 1 → orthogonal; distance 2 → opposite.
    A threshold of 0.7 keeps chunks with cosine similarity ≥ 0.3.
    """
    collection = get_collection()
    count = collection.count()
    if count == 0:
        return None

    results   = collection.query(
        query_texts=[query],
        n_results=min(n_results, count),
    )
    distances = results["distances"][0]
    documents = results["documents"][0]

    relevant = [
        doc for doc, dist in zip(documents, distances)
        if dist <= threshold
    ]
    return relevant if relevant else None


# ── Web search fallback ────────────────────────────────────────────────────────
def search_web(query: str) -> str | None:
    """
    Search DuckDuckGo for query and return formatted snippet text.

    Returns a multi-line string with titles and body excerpts,
    or None if the search fails or returns no results.
    """
    try:
        from ddgs import DDGS
        with DDGS() as ddgs:
            hits = list(ddgs.text(query, max_results=5))
        if not hits:
            return None
        lines = []
        for hit in hits:
            title = hit.get("title", "")
            body  = hit.get("body", "")
            if title or body:
                lines.append(f"**{title}**\n{body}")
        return "\n\n".join(lines) if lines else None
    except Exception:
        return None
